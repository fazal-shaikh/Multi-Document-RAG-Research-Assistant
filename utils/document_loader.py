"""
utils/document_loader.py
Extracts text from various document formats (PDF, DOCX, TXT) and creates LangChain Document objects.
Why it exists: Normalizes different raw strings into a typical Document format ready for ML processing.
What breaks without it: We won't be able to process user-provided knowledge bases, preventing RAG context building.
"""

import logging
from typing import List, Any
from io import BytesIO
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document

# Set up logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

def load_documents(uploaded_files: List[Any]) -> List[Document]:
    """
    Reads multiple uploaded file streams and converts them into LangChain Documents.
    
    Args:
        uploaded_files (List[Any]): Streamlit UploadedFile objects, pre-validated.
        
    Returns:
        List[Document]: List of LangChain document schemas containing text and source metadata.
        
    Example:
        >>> docs = load_documents(validated_files)
    """
    documents = []
    
    # Process each valid file iteratively
    for file in uploaded_files:
        try:
            # Reconstruct byte stream from Streamlit upload object to read it fully
            file_bytes = file.read()
            file_name = file.name
            
            # 1. Handle PDF Documents
            if file_name.endswith(".pdf"):
                reader = PdfReader(BytesIO(file_bytes))
                text_content = ""
                # Extract text iteratively from each page to avoid hitting memory spikes
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                documents.append(Document(page_content=text_content, metadata={"source": file_name}))
                logger.info(f"Loaded PDF: {file_name}")
            
            # 2. Handle DOCX Documents
            elif file_name.endswith(".docx"):
                doc_file = DocxDocument(BytesIO(file_bytes))
                text_content = ""
                # Parse all paragraphs in sequence
                for para in doc_file.paragraphs:
                    text_content += para.text + "\n"
                documents.append(Document(page_content=text_content, metadata={"source": file_name}))
                logger.info(f"Loaded DOCX: {file_name}")
                
            # 3. Handle plain TEXT Documents
            elif file_name.endswith(".txt"):
                # Decode bytes to string assuming UTF-8
                text_content = file_bytes.decode('utf-8')
                documents.append(Document(page_content=text_content, metadata={"source": file_name}))
                logger.info(f"Loaded TXT: {file_name}")
                
            # Fallback (Should be caught by validator naturally)
            else:
                logger.warning(f"File type unsupported at loader level: {file_name}")

        except UnicodeDecodeError as ude:
            # Handle encoding issues in text files gracefully without terminating
            logger.error(f"Cannot decode {file.name}: {ude}")
        except Exception as e:
            # Handle general reading exceptions (e.g., corrupted PDFs)
            logger.error(f"Error reading file {file.name}: {e}")
            
    return documents
